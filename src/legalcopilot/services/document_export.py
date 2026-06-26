"""Document export — convert markdown drafts to downloadable formats.

Currently supports markdown -> DOCX conversion.
"""

import io
import logging

logger = logging.getLogger(__name__)


def markdown_to_docx(markdown_text: str, title: str = "Draft") -> bytes:
    """Convert markdown text to a DOCX file.

    Args:
        markdown_text: The markdown content to convert.
        title: Document title for the DOCX metadata.

    Returns:
        DOCX file bytes.
    """
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Process markdown line by line
    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            text = stripped.split(".", 1)[1].strip() if "." in stripped else stripped
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
